import os
import re

import docx
import docx.enum.style as sty
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Pt
from icecream import ic

# ic.enable()
ic.disable()

syndic = {
    'text': ['宋体', '11', '1.5'],
    'title': ['黑体', '18', '2'],
    'abste': ['仿宋', '10'],
    'absti': ['黑体', '10'],
    'keyword': ['仿宋', '10'],
    'keywordti': ['黑体', '10'],
    '1': ['黑体', '14', '1.5'],
    '1.1': ['黑体', '12','1.5'],
    '1.1.1': ['黑体', '12','1.5'],
    '1.1.1.1': ['黑体', '12'],
    'reft': ['黑体', '10'],
    'ref': ['宋体', '10']
}  # 例子


def change_format(doc_path, syndic):
    doc = docx.Document(doc_path)
    paragraphs = doc.paragraphs

    for key, value in syndic.items():
        custom_style = doc.styles.add_style(key, sty.WD_STYLE_TYPE.PARAGRAPH)
        custom_styleC = doc.styles.add_style(key + 'C', docx.enum.style.WD_STYLE_TYPE.CHARACTER)
        custom_style.font.name = value[0]
        custom_styleC.font.name = value[0]
        custom_style.font.size = Pt(int(value[1]))
        custom_styleC.font.size = Pt(int(value[1]))
        custom_style._element.rPr.rFonts.set(qn('w:eastAsia'), value[0])
        custom_styleC._element.rPr.rFonts.set(qn('w:eastAsia'), value[0])

        if len(value) >= 3:
            custom_style.paragraph_format.line_spacing = Pt(int(value[1]) * float(value[2]))

    # 将全部格式变为syndic中key为text的格式
    for paragraph in paragraphs:
        paragraph.style = 'text'
        match = re.match(r'^(\d+(\.\d+)*)', paragraph.text)
        if not match:
            paragraph.paragraph_format.first_line_indent = 304800

    # 将第一行格式变为syndic中key为title的格式

    first_paragraph = paragraphs[0]
    first_paragraph.style = 'title'
    first_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    sec_paragraph = paragraphs[1]
    sec_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # 从上往下读，将第一个遇到的“摘要：”的所在自然段格式变为syndic中key为abste的格式
    # 将“摘要：”三个字格式变为syndic中key为absti的格式
    for paragraph in paragraphs:
        if "摘要：" in paragraph.text:
            paragraph.style = 'abste'
            paragraph.paragraph_format.first_line_indent = None
            for run in paragraph.runs[:1]:
                run.style = 'abstiC'
            break

    # 从上往下读，将第一个遇到的“关键词：”的所在自然段格式变为syndic中key为keyword的格式
    for paragraph in paragraphs:
        if "关键词：" in paragraph.text:
            paragraph.style = 'keyword'
            paragraph.paragraph_format.first_line_indent = None
            for run in paragraph.runs[:1]:
                run.style = 'keywordtiC'
            break

    # 将所有的“(数字).”开头的行格式变为syndic中key为“1”的格式
    # 将所有的“(数字).(数字)”开头的行格式变为syndic中key为“1.1”的格式
    # 将所有的“(数字).(数字).(数字)”开头的行格式变为syndic中key为“1.1.1”的格式
    # 将所有的“(数字).(数字).(数字).(数字)”开头的行格式变为syndic中key为“1.1.1.1”的格式
    for paragraph in paragraphs:
        text = paragraph.text.strip()
        match = re.match(r'^(\d+(\.\d+)*)', text)  # 使用正则表达式匹配以数字开头的行
        if match:
            line = match.group(1)  # 提取匹配到的数字序列
            levels = line.split('.')
            key_levels = ['1' if level.strip() else '' for level in levels]
            key_format = '.'.join(key_levels)
            ic(key_format)
            paragraph.style = key_format
            if key_format == '1':
                paragraph.add_run("\n")

    # 找到其下一行是以“[”开头的“参考文献”四个字，将其格式变为syndic中key为“reft”的格式
    reference_title_found = False
    for i in range(len(paragraphs) - 1):
        if "参考文献" in paragraphs[i].text and paragraphs[i + 1].text.startswith('['):

            reference_title_found = True
            paragraphs[i].style = 'reft'
            paragraphs[i].paragraph_format.first_line_indent = None
            break

    # 将上一条中“参考文献”以下的，并且以“[”开头的行格式变为syndic中key为“ref”的格式
    reference_found = False
    for i in range(len(paragraphs)):
        if reference_title_found and i > 0 and paragraphs[i].text.startswith('['):
            reference_found = True
            paragraphs[i].style = 'ref'
            paragraphs[i].paragraph_format.first_line_indent = None

    # 将上上一条中“参考文献”以下的，并且以“[”开头的行头三个字符变为“[i]”，i是“参考文献”之下的行数，1开始
    if reference_found:
        reference_count = 1
        for i in range(len(paragraphs)):
            if reference_title_found and i > 0 and paragraphs[i].text.startswith('['):
                paragraphs[i].text = '[{}]{}'.format(reference_count, paragraphs[i].text[3:])
                reference_count += 1
                paragraphs[i].paragraph_format.first_line_indent = None

    # 输出更改格式后的doc
    file_path = 'formatted_doc.docx'
    if os.path.exists(file_path):
        os.remove(file_path)
    doc.save(file_path)


doc_path = '修改格式示例.docx'
change_format(doc_path, syndic)
